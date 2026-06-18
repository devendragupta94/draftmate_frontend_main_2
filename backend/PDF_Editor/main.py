import streamlit as st
import PyPDF2
from PyPDF2 import PdfReader, PdfWriter, PdfMerger
import io
import os
import tempfile
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageDraw, ImageFont
import base64
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table

# Page configuration
st.set_page_config(
    page_title="PDF Toolkit Pro",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for sleek black theme with BLACK text on buttons
st.markdown("""
    <style>
    /* Import Google Font */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
    
    * {
        font-family: 'Inter', sans-serif;
    }
    
    /* Main background - pure black */
    .main {
        background-color: #000000;
    }
    
    /* Content area - dark with subtle border */
    .block-container {
        padding: 2rem;
        background-color: #0a0a0a;
        border-radius: 20px;
        margin: 1rem;
        border: 1px solid #1a1a1a;
        box-shadow: 0 8px 32px rgba(0, 255, 136, 0.1);
    }
    
    /* Header styling with neon accent */
    h1 {
        color: #00ff88;
        font-weight: 700;
        font-size: 3rem !important;
        margin-bottom: 0.5rem;
        text-shadow: 0 0 20px rgba(0, 255, 136, 0.5);
        letter-spacing: -1px;
    }
    
    h2 {
        color: #ffffff;
        font-weight: 600;
        margin-top: 2rem;
        font-size: 1.8rem !important;
    }
    
    h3 {
        color: #00ff88;
        font-weight: 500;
        font-size: 1.3rem !important;
    }
    
    /* Paragraph text - IMPROVED READABILITY */
    p, .stMarkdown {
        color: #d0d0d0 !important;
        font-size: 1.05rem;
    }
    
    /* Better contrast for all text */
    .stMarkdown p, .stMarkdown li, .stMarkdown span {
        color: #d0d0d0 !important;
    }
    
    /* Sidebar styling - pure black with accent */
    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, #000000 0%, #0a0a0a 100%);
        border-right: 2px solid #00ff88;
    }
    
    [data-testid="stSidebar"] h2 {
        color: #00ff88 !important;
        font-size: 1.8rem !important;
        font-weight: 700;
        text-shadow: 0 0 15px rgba(0, 255, 136, 0.5);
    }
    
    [data-testid="stSidebar"] p {
        color: #d0d0d0 !important;
        font-size: 1rem;
    }
    
    /* Radio buttons in sidebar - larger text */
    [data-testid="stSidebar"] .row-widget.stRadio > div {
        background-color: transparent;
    }
    
    [data-testid="stSidebar"] .row-widget.stRadio label {
        color: #ffffff !important;
        font-size: 1.1rem !important;
        font-weight: 500;
        padding: 0.8rem 1rem;
        border-radius: 10px;
        transition: all 0.3s ease;
        cursor: pointer;
        display: block;
        margin: 0.3rem 0;
    }
    
    [data-testid="stSidebar"] .row-widget.stRadio label:hover {
        background-color: #1a1a1a;
        color: #00ff88 !important;
        transform: translateX(5px);
    }
    
    [data-testid="stSidebar"] .row-widget.stRadio input:checked + label {
        background: linear-gradient(135deg, #00ff88 0%, #00cc6a 100%);
        color: #000000 !important;
        font-weight: 600;
        box-shadow: 0 4px 15px rgba(0, 255, 136, 0.3);
    }
    
    /* PRIMARY BUTTON STYLING - BLACK TEXT ON GREEN */
    .stButton > button {
        background: linear-gradient(135deg, #00ff88 0%, #00cc6a 100%) !important;
        color: #000000 !important;
        border: none !important;
        border-radius: 12px !important;
        padding: 0.75rem 2rem !important;
        font-weight: 700 !important;
        font-size: 1.1rem !important;
        transition: all 0.3s ease !important;
        box-shadow: 0 4px 15px rgba(0, 255, 136, 0.3) !important;
        text-transform: uppercase !important;
        letter-spacing: 0.5px !important;
    }
    
    .stButton > button:hover {
        transform: translateY(-3px) !important;
        box-shadow: 0 6px 25px rgba(0, 255, 136, 0.5) !important;
        background: linear-gradient(135deg, #00ffaa 0%, #00dd7a 100%) !important;
        color: #000000 !important;
    }
    
    .stButton > button:active {
        transform: translateY(-1px) !important;
    }
    
    .stButton > button p {
        color: #000000 !important;
    }
    
    .stButton > button span {
        color: #000000 !important;
    }
    
    .stButton > button div {
        color: #000000 !important;
    }
    
    /* DOWNLOAD BUTTON STYLING - BLACK TEXT ON GREEN */
    .stDownloadButton > button {
        background: linear-gradient(135deg, #00ff88 0%, #00cc6a 100%) !important;
        color: #000000 !important;
        border: none !important;
        border-radius: 12px !important;
        padding: 0.75rem 2rem !important;
        font-weight: 700 !important;
        font-size: 1rem !important;
        box-shadow: 0 4px 15px rgba(0, 255, 136, 0.3) !important;
    }
    
    .stDownloadButton > button:hover {
        transform: translateY(-2px) !important;
        box-shadow: 0 6px 20px rgba(0, 255, 136, 0.5) !important;
        color: #000000 !important;
    }
    
    .stDownloadButton > button p {
        color: #000000 !important;
    }
    
    .stDownloadButton > button span {
        color: #000000 !important;
    }
    
    .stDownloadButton > button div {
        color: #000000 !important;
    }
    
    /* File uploader styling */
    [data-testid="stFileUploader"] {
        background: #1a1a1a;
        border-radius: 15px;
        padding: 2rem;
        border: 2px dashed #00ff88;
        transition: all 0.3s ease;
    }
    
    [data-testid="stFileUploader"]:hover {
        border-color: #00ffaa;
        background: #222222;
        box-shadow: 0 0 20px rgba(0, 255, 136, 0.2);
    }
    
    [data-testid="stFileUploader"] label {
        color: #ffffff !important;
        font-size: 1.1rem !important;
    }
    
    [data-testid="stFileUploader"] small {
        color: #a0a0a0 !important;
    }
    
    /* Info/Warning/Success boxes - IMPROVED READABILITY */
    .stAlert {
        border-radius: 12px;
        border-left: 4px solid #00ff88;
        background-color: #1a1a1a;
        color: #e0e0e0 !important;
    }
    
    [data-baseweb="notification"] {
        background-color: #1a1a1a;
        border-left: 4px solid #00ff88;
    }
    
    /* Success message */
    .stSuccess {
        background-color: #0d2818 !important;
        border-left: 4px solid #00ff88 !important;
        color: #00ff88 !important;
    }
    
    .stSuccess p, .stSuccess div {
        color: #00ff88 !important;
    }
    
    /* Error message */
    .stError {
        background-color: #2a0d0d !important;
        border-left: 4px solid #ff4444 !important;
        color: #ff8888 !important;
    }
    
    .stError p, .stError div {
        color: #ff8888 !important;
    }
    
    /* Warning message */
    .stWarning {
        background-color: #2a2410 !important;
        border-left: 4px solid #ffaa00 !important;
        color: #ffdd88 !important;
    }
    
    .stWarning p, .stWarning div {
        color: #ffdd88 !important;
    }
    
    /* Info message */
    .stInfo {
        background-color: #0d1f2a !important;
        border-left: 4px solid #00aaff !important;
        color: #88ddff !important;
    }
    
    .stInfo p, .stInfo div {
        color: #88ddff !important;
    }
    
    /* Metrics */
    [data-testid="stMetricValue"] {
        color: #00ff88 !important;
        font-size: 2.5rem !important;
        font-weight: 700 !important;
        text-shadow: 0 0 15px rgba(0, 255, 136, 0.3);
    }
    
    [data-testid="stMetricLabel"] {
        color: #c0c0c0 !important;
        font-size: 1rem !important;
        font-weight: 500 !important;
    }
    
    [data-testid="stMetricDelta"] {
        color: #00ff88 !important;
    }
    
    /* PDF Viewer container */
    .pdf-viewer {
        border: 2px solid #00ff88;
        border-radius: 15px;
        padding: 1.5rem;
        background: #0a0a0a;
        max-height: 600px;
        overflow-y: auto;
        box-shadow: 0 4px 20px rgba(0, 255, 136, 0.2);
    }
    
    /* Scrollbar styling */
    ::-webkit-scrollbar {
        width: 10px;
        height: 10px;
    }
    
    ::-webkit-scrollbar-track {
        background: #0a0a0a;
    }
    
    ::-webkit-scrollbar-thumb {
        background: #00ff88;
        border-radius: 5px;
    }
    
    ::-webkit-scrollbar-thumb:hover {
        background: #00cc6a;
    }
    
    /* Slider styling */
    .stSlider {
        padding: 1rem 0;
    }
    
    .stSlider > div > div > div {
        background-color: #1a1a1a;
    }
    
    .stSlider > div > div > div > div {
        background-color: #00ff88;
    }
    
    /* Text input - IMPROVED READABILITY */
    .stTextInput > div > div > input {
        background-color: #1a1a1a !important;
        color: #ffffff !important;
        border: 2px solid #333333 !important;
        border-radius: 10px;
        font-size: 1rem;
        padding: 0.75rem;
    }
    
    .stTextInput > div > div > input:focus {
        border-color: #00ff88 !important;
        box-shadow: 0 0 15px rgba(0, 255, 136, 0.3);
    }
    
    .stTextInput label {
        color: #e0e0e0 !important;
        font-size: 1rem !important;
        font-weight: 500 !important;
    }
    
    /* Select box - IMPROVED READABILITY */
    .stSelectbox > div > div > div {
        background-color: #1a1a1a !important;
        color: #ffffff !important;
        border: 2px solid #333333 !important;
        border-radius: 10px;
    }
    
    .stSelectbox label {
        color: #e0e0e0 !important;
        font-size: 1rem !important;
        font-weight: 500 !important;
    }
    
    /* Number input - IMPROVED READABILITY */
    .stNumberInput label {
        color: #e0e0e0 !important;
        font-size: 1rem !important;
        font-weight: 500 !important;
    }
    
    .stNumberInput input {
        background-color: #1a1a1a !important;
        color: #ffffff !important;
        border: 2px solid #333333 !important;
    }
    
    /* Radio buttons (horizontal) */
    .row-widget.stRadio > div {
        flex-direction: row;
        gap: 1rem;
    }
    
    .row-widget.stRadio label {
        background-color: #1a1a1a;
        padding: 0.6rem 1.2rem;
        border-radius: 10px;
        border: 2px solid #333333;
        color: #ffffff !important;
        font-weight: 500;
        transition: all 0.3s ease;
    }
    
    .row-widget.stRadio label:hover {
        border-color: #00ff88;
        background-color: #222222;
    }
    
    .row-widget.stRadio input:checked + label {
        background-color: #00ff88 !important;
        color: #000000 !important;
        border-color: #00ff88;
        font-weight: 600;
    }
    
    /* Expander */
    .streamlit-expanderHeader {
        background-color: #1a1a1a;
        color: #ffffff !important;
        border-radius: 10px;
        border: 1px solid #333333;
        font-size: 1.1rem;
        font-weight: 500;
    }
    
    .streamlit-expanderHeader:hover {
        border-color: #00ff88;
        background-color: #222222;
    }
    
    .streamlit-expanderContent {
        background-color: #0a0a0a;
        border: 1px solid #333333;
        border-top: none;
    }
    
    /* Divider */
    hr {
        border-color: #333333;
        margin: 2rem 0;
    }
    
    /* Column styling */
    [data-testid="column"] {
        background-color: transparent;
    }
    
    /* Spinner */
    .stSpinner > div {
        border-top-color: #00ff88 !important;
    }
    </style>
    """, unsafe_allow_html=True)

# Title with emoji
st.markdown("<h1>📄 PDF Toolkit</h1>", unsafe_allow_html=True)
st.markdown("<p style='color: #d0d0d0; font-size: 1.1rem;'><strong>All-in-one PDF manipulation solution</strong></p>", unsafe_allow_html=True)
st.markdown("---")

# Sidebar for navigation
st.sidebar.markdown("<h2>🛠️ Tools</h2>", unsafe_allow_html=True)
tool = st.sidebar.radio(
    "",
    [
        "🔗 Merge PDFs",
        "✂️ Split PDF",
        "🗜️ Compress PDF",
        "📄➡️📝 PDF to Word",
        "📝➡️📄 Word to PDF",
        "🔄 Rotate PDF",
        "💧 Add Watermark",
        "✏️ Edit PDF"
    ],
    label_visibility="collapsed"
)

st.sidebar.markdown("---")
st.sidebar.markdown("<p style='color: #d0d0d0; font-size: 1rem;'>💡 Upload your files and use the tools to manipulate PDFs easily!</p>", unsafe_allow_html=True)


# Helper Functions
def display_pdf_pages_as_images(pdf_file, page_range=None, max_pages=None):
    """Display PDF pages as images for preview with optional range"""
    try:
        pdf_file.seek(0)
        pdf_bytes = pdf_file.read()
        pdf_file.seek(0)
        
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        total_pages = len(doc)
        
        st.markdown("<div class='pdf-viewer'>", unsafe_allow_html=True)
        
        # Determine which pages to show
        if page_range:
            start_page, end_page = page_range
            pages_to_show = range(start_page - 1, min(end_page, total_pages))
        elif max_pages:
            pages_to_show = range(min(total_pages, max_pages))
        else:
            pages_to_show = range(total_pages)
        
        for page_num in pages_to_show:
            page = doc[page_num]
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            
            st.image(img, caption=f"Page {page_num + 1}", use_container_width=True)
        
        if max_pages and total_pages > max_pages:
            st.info(f"Showing first {max_pages} pages of {total_pages} total pages")
        elif page_range:
            st.info(f"Showing pages {page_range[0]} to {min(page_range[1], total_pages)} of {total_pages} total pages")
        
        st.markdown("</div>", unsafe_allow_html=True)
        doc.close()
        
    except Exception as e:
        st.error(f"Could not display preview: {str(e)}")


def show_preview_options(uploaded_file, label="Preview PDF"):
    """Show preview options with page range selection"""
    with st.expander(f"👁️ {label}"):
        uploaded_file.seek(0)
        pdf_bytes = uploaded_file.read()
        uploaded_file.seek(0)
        
        doc_temp = fitz.open(stream=pdf_bytes, filetype="pdf")
        total_pages = len(doc_temp)
        doc_temp.close()
        
        st.write(f"**Total Pages:** {total_pages}")
        
        preview_option = st.radio(
            "Preview Option:",
            ["First 5 pages", "Custom range", "All pages"],
            horizontal=True,
            key=f"preview_{label}"
        )
        
        if preview_option == "First 5 pages":
            display_pdf_pages_as_images(uploaded_file, max_pages=5)
        elif preview_option == "Custom range":
            col1, col2 = st.columns(2)
            with col1:
                start_page = st.number_input("Start Page", min_value=1, max_value=total_pages, value=1, key=f"start_{label}")
            with col2:
                end_page = st.number_input("End Page", min_value=start_page, max_value=total_pages, value=min(start_page + 4, total_pages), key=f"end_{label}")
            
            if st.button("Show Preview", key=f"show_{label}"):
                display_pdf_pages_as_images(uploaded_file, page_range=(start_page, end_page))
        else:  # All pages
            if total_pages > 20:
                st.warning(f"⚠️ This document has {total_pages} pages. Loading all pages may take time.")
                if st.button("Load All Pages", key=f"load_all_{label}"):
                    display_pdf_pages_as_images(uploaded_file)
            else:
                display_pdf_pages_as_images(uploaded_file)


def merge_pdfs(pdf_files):
    """Merge multiple PDF files"""
    merger = PdfMerger()
    try:
        for pdf_file in pdf_files:
            merger.append(pdf_file)
        
        output = io.BytesIO()
        merger.write(output)
        merger.close()
        output.seek(0)
        return output
    except Exception as e:
        st.error(f"Error merging PDFs: {str(e)}")
        return None


def split_pdf(pdf_file, page_ranges):
    """Split PDF into specified page ranges"""
    try:
        reader = PdfReader(pdf_file)
        outputs = []
        
        for page_range in page_ranges:
            writer = PdfWriter()
            start, end = page_range
            
            for page_num in range(start - 1, min(end, len(reader.pages))):
                writer.add_page(reader.pages[page_num])
            
            output = io.BytesIO()
            writer.write(output)
            output.seek(0)
            outputs.append(output)
        
        return outputs
    except Exception as e:
        st.error(f"Error splitting PDF: {str(e)}")
        return None


def compress_pdf_advanced(pdf_file, target_size_kb=None, compression_level="medium"):
    """Advanced PDF compression with better control"""
    try:
        pdf_bytes = pdf_file.read()
        pdf_file.seek(0)
        original_size_kb = len(pdf_bytes) / 1024
        
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        
        # Compression settings based on level
        if compression_level == "low":
            image_quality = 85
            dpi = 150
        elif compression_level == "medium":
            image_quality = 60
            dpi = 120
        else:  # high
            image_quality = 40
            dpi = 90
        
        # If target size is specified, calculate required quality
        if target_size_kb:
            size_ratio = target_size_kb / original_size_kb
            if size_ratio < 0.3:
                image_quality = 30
                dpi = 72
            elif size_ratio < 0.5:
                image_quality = 50
                dpi = 96
            elif size_ratio < 0.7:
                image_quality = 65
                dpi = 120
        
        # Process each page
        for page_num in range(len(doc)):
            page = doc[page_num]
            image_list = page.get_images(full=True)
            
            for img_index, img in enumerate(image_list):
                xref = img[0]
                
                try:
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image = Image.open(io.BytesIO(image_bytes))
                    
                    width, height = image.size
                    new_width = int(width * (dpi / 150))
                    new_height = int(height * (dpi / 150))
                    
                    if new_width > 0 and new_height > 0:
                        image = image.resize((new_width, new_height), Image.Resampling.LANCZOS)
                    
                    if image.mode != 'RGB':
                        image = image.convert('RGB')
                    
                    img_output = io.BytesIO()
                    image.save(img_output, format='JPEG', quality=image_quality, optimize=True)
                    img_bytes = img_output.getvalue()
                    
                    img_output.seek(0)
                    page.replace_image(xref, stream=img_bytes)
                    
                except Exception as img_error:
                    continue
        
        output = io.BytesIO()
        # Use less aggressive garbage collection to prevent corruption
        doc.save(
            output,
            garbage=3, 
            deflate=True,
            clean=True,
            deflate_images=True,
            deflate_fonts=True
        )
        doc.close()
        
        # Validate output
        if not validate_pdf(output.getvalue()):
            st.warning("Compression produced invalid file, returning original.")
            pdf_file.seek(0)
            return pdf_file
        
        output.seek(0)
        final_size_kb = len(output.getvalue()) / 1024
        
        if target_size_kb and final_size_kb > target_size_kb * 1.1:
            output.seek(0)
            doc2 = fitz.open(stream=output.getvalue(), filetype="pdf")
            
            for page_num in range(len(doc2)):
                page = doc2[page_num]
                image_list = page.get_images(full=True)
                
                for img_index, img in enumerate(image_list):
                    xref = img[0]
                    try:
                        base_image = doc2.extract_image(xref)
                        image_bytes = base_image["image"]
                        image = Image.open(io.BytesIO(image_bytes))
                        
                        width, height = image.size
                        image = image.resize((int(width * 0.7), int(height * 0.7)), Image.Resampling.LANCZOS)
                        
                        if image.mode != 'RGB':
                            image = image.convert('RGB')
                        
                        img_output = io.BytesIO()
                        image.save(img_output, format='JPEG', quality=max(20, image_quality - 20), optimize=True)
                        img_bytes = img_output.getvalue()
                        
                        img_output.seek(0)
                        page.replace_image(xref, stream=img_bytes)
                    except:
                        continue
            
            output2 = io.BytesIO()
            doc2.save(output2, garbage=4, deflate=True, clean=True)
            doc2.close()
            output2.seek(0)
            return output2
        
        return output
        
    except Exception as e:
        st.error(f"Error compressing PDF: {str(e)}")
        return None


def validate_pdf(pdf_bytes):
    """Validate if PDF is valid and not empty"""
    try:
        if len(pdf_bytes) < 100:  # Too small to be valid
            return False
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        if len(doc) == 0:
            return False
        doc.close()
        return True
    except:
        return False


def pdf_to_word_pymupdf(pdf_file):
    """Convert PDF to Word document using PyMuPDF"""
    try:
        pdf_bytes = pdf_file.read()
        pdf_file.seek(0)
        
        pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
        doc = Document()
        
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            text = page.get_text()
            
            if page_num > 0:
                doc.add_page_break()
            
            heading = doc.add_paragraph(f"Page {page_num + 1}")
            heading.style = 'Heading 2'
            
            if text.strip():
                paragraphs = text.split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        para = doc.add_paragraph(para_text.strip())
                        para.style = 'Normal'
            
            try:
                image_list = page.get_images()
                for img_index, img in enumerate(image_list):
                    xref = img[0]
                    base_image = pdf_document.extract_image(xref)
                    image_bytes = base_image["image"]
                    
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_img:
                        tmp_img.write(image_bytes)
                        tmp_img_path = tmp_img.name
                    
                    try:
                        doc.add_picture(tmp_img_path, width=Inches(4))
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except:
                        pass
                    finally:
                        try:
                            os.unlink(tmp_img_path)
                        except:
                            pass
            except:
                pass
        
        pdf_document.close()
        
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        return output
    except Exception as e:
        st.error(f"Error converting PDF to Word: {str(e)}")
        return None


def word_to_pdf_improved(docx_file):
    """Improved Word to PDF conversion using reportlab"""
    try:
        docx_file.seek(0)
        doc = Document(docx_file)
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_pdf:
            tmp_pdf_path = tmp_pdf.name
        
        pdf_doc = SimpleDocTemplate(tmp_pdf_path, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        for para in doc.paragraphs:
            if para.text.strip():
                if para.style.name.startswith('Heading'):
                    style = styles['Heading1']
                else:
                    style = styles['Normal']
                
                p = Paragraph(para.text, style)
                story.append(p)
                story.append(Spacer(1, 12))
        
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text)
                table_data.append(row_data)
            
            if table_data:
                t = Table(table_data)
                story.append(t)
                story.append(Spacer(1, 12))
        
        pdf_doc.build(story)
        
        with open(tmp_pdf_path, 'rb') as f:
            output = io.BytesIO(f.read())
        
        os.unlink(tmp_pdf_path)
        
        output.seek(0)
        return output
    except Exception as e:
        st.error(f"Error converting Word to PDF: {str(e)}")
        return None


def add_watermark_simple(pdf_file, watermark_text=None, watermark_image=None, opacity=0.3, position="center", font_size=50, angle=45, scale=1.0):
    """Add watermark to PDF - SIMPLIFIED VERSION THAT WORKS"""
    try:
        pdf_bytes = pdf_file.read()
        pdf_file.seek(0)
        
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            page_width = page.rect.width
            page_height = page.rect.height
            
            if watermark_text:
                # Create a PIL image for the watermark
                text_length = len(watermark_text)
                img_width = int(text_length * font_size * 0.7)
                img_height = int(font_size * 2)
                
                # Create transparent image
                watermark_img = Image.new('RGBA', (img_width, img_height), (255, 255, 255, 0))
                draw = ImageDraw.Draw(watermark_img)
                
                # Try to use a font, fallback to default
                try:
                    font = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", font_size)
                except:
                    try:
                        font = ImageFont.truetype("/Library/Fonts/Arial.ttf", font_size)
                    except:
                        font = ImageFont.load_default()
                
                # Calculate text color with opacity
                alpha = int(opacity * 255)
                text_color = (128, 128, 128, alpha)  # Gray with alpha
                
                # Draw text
                draw.text((10, 10), watermark_text, fill=text_color, font=font)
                
                # Rotate the image
                watermark_img = watermark_img.rotate(angle, expand=True)
                
                # Save to bytes
                img_byte_arr = io.BytesIO()
                watermark_img.save(img_byte_arr, format='PNG')
                img_bytes = img_byte_arr.getvalue()
                
                # Calculate position
                if position == "center":
                    x = (page_width - watermark_img.width) / 2
                    y = (page_height - watermark_img.height) / 2
                elif position == "top-left":
                    x, y = 50, page_height - watermark_img.height - 50
                elif position == "top-right":
                    x = page_width - watermark_img.width - 50
                    y = page_height - watermark_img.height - 50
                elif position == "bottom-left":
                    x, y = 50, 50
                elif position == "bottom-right":
                    x = page_width - watermark_img.width - 50
                    y = 50
                else:
                    x = (page_width - watermark_img.width) / 2
                    y = (page_height - watermark_img.height) / 2
                
                # Insert image
                img_rect = fitz.Rect(x, y, x + watermark_img.width, y + watermark_img.height)
                page.insert_image(img_rect, stream=img_bytes, overlay=True)
                
            elif watermark_image:
                # Calculate position for image
                base_width = 200
                base_height = 200
                img_width = int(base_width * scale)
                img_height = int(base_height * scale)
                
                if position == "center":
                    x0, y0 = (page_width - img_width) / 2, (page_height - img_height) / 2
                elif position == "top-left":
                    x0, y0 = 50, page_height - img_height - 50
                elif position == "top-right":
                    x0, y0 = page_width - img_width - 50, page_height - img_height - 50
                elif position == "bottom-left":
                    x0, y0 = 50, 50
                elif position == "bottom-right":
                    x0, y0 = page_width - img_width - 50, 50
                else:
                    x0, y0 = (page_width - img_width) / 2, (page_height - img_height) / 2
                
                # Open image and add transparency
                img = Image.open(io.BytesIO(watermark_image))
                
                # Convert to RGBA if not already
                if img.mode != 'RGBA':
                    img = img.convert('RGBA')
                
                # Adjust opacity
                alpha = img.split()[3]
                alpha = alpha.point(lambda p: int(p * opacity))
                img.putalpha(alpha)
                
                # Resize image
                img = img.resize((img_width, img_height), Image.Resampling.LANCZOS)
                
                # Rotate image
                img = img.rotate(angle, expand=True)
                
                # Update dimensions after rotation
                img_width, img_height = img.size
                
                # Save to bytes
                img_byte_arr = io.BytesIO()
                img.save(img_byte_arr, format='PNG')
                img_bytes = img_byte_arr.getvalue()
                
                img_rect = fitz.Rect(x0, y0, x0 + img_width, y0 + img_height)
                page.insert_image(img_rect, stream=img_bytes, overlay=True)
        
        output = io.BytesIO()
        doc.save(output)
        doc.close()
        
        output.seek(0)
        return output
    except Exception as e:
        st.error(f"Error adding watermark: {str(e)}")
        import traceback
        st.error(traceback.format_exc())
        return None


def rotate_pdf(pdf_file, rotation_angle, pages_to_rotate):
    """Rotate specified pages of PDF"""
    try:
        reader = PdfReader(pdf_file)
        writer = PdfWriter()
        
        for i, page in enumerate(reader.pages):
            if pages_to_rotate == "all" or i + 1 in pages_to_rotate:
                page.rotate(rotation_angle)
            writer.add_page(page)
        
        output = io.BytesIO()
        writer.write(output)
        output.seek(0)
        return output
    except Exception as e:
        st.error(f"Error rotating PDF: {str(e)}")
        return None


def extract_pages(pdf_file, pages_to_keep):
    """Extract specific pages from PDF"""
    try:
        reader = PdfReader(pdf_file)
        writer = PdfWriter()
        
        for page_num in pages_to_keep:
            if 0 < page_num <= len(reader.pages):
                writer.add_page(reader.pages[page_num - 1])
        
        output = io.BytesIO()
        writer.write(output)
        output.seek(0)
        return output
    except Exception as e:
        st.error(f"Error extracting pages: {str(e)}")
        return None


# Tool Implementations
if tool == "🔗 Merge PDFs":
    st.header("🔗 Merge Multiple PDFs")
    st.write("Upload multiple PDF files to merge them into one document.")
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        uploaded_files = st.file_uploader(
            "Choose PDF files",
            type=['pdf'],
            accept_multiple_files=True,
            key="merge"
        )
    
    if uploaded_files:
        with col2:
            st.metric("Files Uploaded", len(uploaded_files))
        
        st.success(f"**{len(uploaded_files)} file(s) ready to merge:**")
        for i, file in enumerate(uploaded_files, 1):
            st.write(f"  {i}. 📄 {file.name}")
        
        if st.button("🔗 Merge PDFs", type="primary", use_container_width=True):
            with st.spinner("Merging PDFs..."):
                merged_pdf = merge_pdfs(uploaded_files)
                
                if merged_pdf:
                    st.success("✅ PDFs merged successfully!")
                    
                    st.download_button(
                        label="📥 Download Merged PDF",
                        data=merged_pdf,
                        file_name="merged_document.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
                    
                    show_preview_options(merged_pdf, "Merged PDF Result")


elif tool == "✂️ Split PDF":
    st.header("✂️ Split PDF")
    st.write("Split a PDF file into multiple documents based on page ranges.")
    
    uploaded_file = st.file_uploader("Choose a PDF file", type=['pdf'], key="split")
    
    if uploaded_file:
        reader = PdfReader(uploaded_file)
        total_pages = len(reader.pages)
        
        col1, col2 = st.columns(2)
        with col1:
            st.metric("Total Pages", total_pages)
        
        show_preview_options(uploaded_file, "Original PDF")
        
        split_option = st.radio(
            "Split option:",
            ["Split by page ranges", "Extract individual pages"],
            horizontal=True
        )
        
        if split_option == "Split by page ranges":
            st.info("💡 Enter page ranges (e.g., 1-3, 4-6, 7-10)")
            ranges_input = st.text_input(
                "Page ranges (comma-separated):",
                placeholder="1-3, 4-6, 7-10"
            )
            
            if st.button("✂️ Split PDF", type="primary", use_container_width=True) and ranges_input:
                try:
                    ranges = []
                    for r in ranges_input.split(','):
                        r = r.strip()
                        if '-' in r:
                            start, end = map(int, r.split('-'))
                            ranges.append((start, end))
                        else:
                            page = int(r)
                            ranges.append((page, page))
                    
                    with st.spinner("Splitting PDF..."):
                        uploaded_file.seek(0)
                        split_pdfs = split_pdf(uploaded_file, ranges)
                        
                        if split_pdfs:
                            st.success(f"✅ PDF split into {len(split_pdfs)} document(s)!")
                            
                            cols = st.columns(min(len(split_pdfs), 3))
                            for i, pdf in enumerate(split_pdfs):
                                with cols[i % 3]:
                                    st.download_button(
                                        label=f"📥 Part {i+1}",
                                        data=pdf,
                                        file_name=f"split_part_{i+1}.pdf",
                                        mime="application/pdf",
                                        key=f"split_{i}",
                                        use_container_width=True
                                    )
                except Exception as e:
                    st.error(f"Invalid page range format: {str(e)}")
        
        else:
            st.info("💡 Enter page numbers separated by commas")
            pages_input = st.text_input(
                "Pages to extract:",
                placeholder="1, 3, 5, 7"
            )
            
            if st.button("✂️ Extract Pages", type="primary", use_container_width=True) and pages_input:
                try:
                    pages = [int(p.strip()) for p in pages_input.split(',')]
                    
                    with st.spinner("Extracting pages..."):
                        uploaded_file.seek(0)
                        extracted_pdf = extract_pages(uploaded_file, pages)
                        
                        if extracted_pdf:
                            st.success("✅ Pages extracted successfully!")
                            st.download_button(
                                label="📥 Download Extracted Pages",
                                data=extracted_pdf,
                                file_name="extracted_pages.pdf",
                                mime="application/pdf",
                                use_container_width=True
                            )
                            
                            show_preview_options(extracted_pdf, "Extracted Pages Result")
                except Exception as e:
                    st.error(f"Invalid page numbers: {str(e)}")


elif tool == "🗜️ Compress PDF":
    st.header("🗜️ Compress PDF")
    st.write("Reduce the file size of your PDF document with customizable compression.")
    
    uploaded_file = st.file_uploader("Choose a PDF file", type=['pdf'], key="compress")
    
    if uploaded_file:
        original_size = len(uploaded_file.getvalue()) / 1024
        
        col1, col2 = st.columns(2)
        with col1:
            st.metric("Original Size", f"{original_size:.2f} KB")
        
        st.subheader("Compression Settings")
        
        compression_mode = st.radio(
            "Compression Mode:",
            ["Compression Level", "Target Size"],
            horizontal=True
        )
        
        target_size = None
        compression_level = "medium"
        
        if compression_mode == "Target Size":
            target_size = st.slider(
                "Target Size (KB)",
                min_value=int(original_size * 0.1),
                max_value=int(original_size * 0.9),
                value=int(original_size * 0.5),
                help="The tool will try to compress the PDF to approximately this size"
            )
            st.info(f"🎯 Target: {target_size} KB (Reduction: {((original_size - target_size) / original_size * 100):.1f}%)")
        else:
            compression_level = st.select_slider(
                "Compression Level",
                options=["low", "medium", "high"],
                value="medium",
                help="Low = Better quality, larger size | High = Lower quality, smaller size"
            )
            
            if compression_level == "low":
                st.info("📊 Low compression: ~15-25% size reduction, maintains good quality")
            elif compression_level == "medium":
                st.info("📊 Medium compression: ~40-50% size reduction, balanced quality")
            else:
                st.info("📊 High compression: ~60-70% size reduction, lower quality")
        
        if st.button("🗜️ Compress PDF", type="primary", use_container_width=True):
            with st.spinner("Compressing PDF... This may take a moment..."):
                uploaded_file.seek(0)
                compressed_pdf = compress_pdf_advanced(
                    uploaded_file, 
                    target_size_kb=target_size, 
                    compression_level=compression_level
                )
                
                if compressed_pdf:
                    compressed_size = len(compressed_pdf.getvalue()) / 1024
                    reduction = ((original_size - compressed_size) / original_size) * 100
                    
                    st.success("✅ PDF compressed successfully!")
                    
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.metric("Original Size", f"{original_size:.2f} KB")
                    with col2:
                        st.metric("Compressed Size", f"{compressed_size:.2f} KB", f"-{reduction:.1f}%")
                    with col3:
                        st.metric("Space Saved", f"{original_size - compressed_size:.2f} KB")
                    
                    if target_size and compressed_size > target_size * 1.2:
                        st.warning(f"⚠️ Could not reach target size. Achieved {compressed_size:.2f} KB (Target was {target_size} KB)")
                    
                    st.download_button(
                        label="📥 Download Compressed PDF",
                        data=compressed_pdf,
                        file_name="compressed_document.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
                    
                    show_preview_options(compressed_pdf, "Compressed PDF Result")


elif tool == "📄➡️📝 PDF to Word":
    st.header("📄➡️📝 Convert PDF to Word")
    st.write("Convert your PDF document to an editable Word file (.docx).")
    
    uploaded_file = st.file_uploader("Choose a PDF file", type=['pdf'], key="pdf2word")
    
    if uploaded_file:
        col1, col2 = st.columns([2, 1])
        
        with col1:
            st.info("ℹ️ This process may take a few moments depending on the file size.")
            st.warning("⚠️ Complex formatting may not be preserved perfectly.")
        
        show_preview_options(uploaded_file, "Original PDF")
        
        if st.button("🔄 Convert to Word", type="primary", use_container_width=True):
            with st.spinner("Converting PDF to Word..."):
                uploaded_file.seek(0)
                word_doc = pdf_to_word_pymupdf(uploaded_file)
                
                if word_doc:
                    st.success("✅ Conversion successful!")
                    st.balloons()
                    st.download_button(
                        label="📥 Download Word Document",
                        data=word_doc,
                        file_name="converted_document.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )


elif tool == "📝➡️📄 Word to PDF":
    st.header("📝➡️📄 Convert Word to PDF")
    st.write("Convert your Word document (.docx) to PDF format.")
    
    uploaded_file = st.file_uploader("Choose a Word file", type=['docx'], key="word2pdf")
    
    if uploaded_file:
        st.info("ℹ️ Your document will be converted to PDF format with preserved formatting.")
        
        if st.button("🔄 Convert to PDF", type="primary", use_container_width=True):
            with st.spinner("Converting Word to PDF..."):
                uploaded_file.seek(0)
                pdf_doc = word_to_pdf_improved(uploaded_file)
                
                if pdf_doc:
                    st.success("✅ Conversion successful!")
                    st.balloons()
                    
                    st.download_button(
                        label="📥 Download PDF Document",
                        data=pdf_doc,
                        file_name="converted_document.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
                    
                    show_preview_options(pdf_doc, "Converted PDF")


elif tool == "🔄 Rotate PDF":
    st.header("🔄 Rotate PDF Pages")
    st.write("Rotate pages in your PDF document.")
    
    uploaded_file = st.file_uploader("Choose a PDF file", type=['pdf'], key="rotate")
    
    if uploaded_file:
        reader = PdfReader(uploaded_file)
        total_pages = len(reader.pages)
        
        st.metric("Total Pages", total_pages)
        
        show_preview_options(uploaded_file, "Original PDF")
        
        col1, col2 = st.columns(2)
        
        with col1:
            rotation_angle = st.selectbox(
                "Rotation angle:",
                [90, 180, 270],
                format_func=lambda x: f"{x}° clockwise"
            )
        
        with col2:
            rotate_option = st.radio(
                "Pages to rotate:",
                ["All pages", "Specific pages"],
                horizontal=True
            )
        
        pages_to_rotate = "all"
        if rotate_option == "Specific pages":
            pages_input = st.text_input(
                "Enter page numbers (comma-separated):",
                placeholder="1, 3, 5"
            )
            if pages_input:
                try:
                    pages_to_rotate = [int(p.strip()) for p in pages_input.split(',')]
                except:
                    st.error("Invalid page numbers format")
                    pages_to_rotate = []
        
        if st.button("🔄 Rotate PDF", type="primary", use_container_width=True):
            with st.spinner("Rotating PDF..."):
                uploaded_file.seek(0)
                rotated_pdf = rotate_pdf(uploaded_file, rotation_angle, pages_to_rotate)
                
                if rotated_pdf:
                    st.success("✅ PDF rotated successfully!")
                    
                    st.download_button(
                        label="📥 Download Rotated PDF",
                        data=rotated_pdf,
                        file_name="rotated_document.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
                    
                    show_preview_options(rotated_pdf, "Rotated PDF Result")


elif tool == "💧 Add Watermark":
    st.header("💧 Add Watermark to PDF")
    st.write("Add text or image watermark to your PDF document.")
    
    uploaded_file = st.file_uploader("Choose a PDF file", type=['pdf'], key="watermark")
    
    if uploaded_file:
        show_preview_options(uploaded_file, "Original PDF")
        
        st.subheader("Watermark Settings")
        
        watermark_type = st.radio(
            "Watermark Type:",
            ["Text", "Image"],
            horizontal=True
        )
        
        watermark_text = None
        watermark_image = None
        
        col1, col2 = st.columns(2)
        
        with col1:
            if watermark_type == "Text":
                watermark_text = st.text_input(
                    "Watermark Text:",
                    value="CONFIDENTIAL",
                    placeholder="Enter watermark text"
                )
                font_size = st.slider("Font Size", 20, 100, 50)
            else:
                uploaded_watermark = st.file_uploader(
                    "Upload Watermark Image",
                    type=['png', 'jpg', 'jpeg'],
                    key="watermark_img"
                )
                if uploaded_watermark:
                    watermark_image = uploaded_watermark.read()
                    st.image(watermark_image, caption="Watermark Preview", width=200)
                font_size = 50
        
        with col2:
            opacity = st.slider(
                "Opacity",
                min_value=0.1,
                max_value=1.0,
                value=0.3,
                step=0.1,
                help="Lower values = more transparent"
            )
            
            angle = st.slider("Rotation Angle", -180, 180, 45, step=5)
            scale = st.slider("Watermark Scale", 0.1, 3.0, 1.0, step=0.1)
            
            position = st.selectbox(
                "Position:",
                ["center", "top-left", "top-right", "bottom-left", "bottom-right"]
            )
        
        if st.button("💧 Add Watermark", type="primary", use_container_width=True):
            if watermark_type == "Text" and not watermark_text:
                st.error("Please enter watermark text")
            elif watermark_type == "Image" and not watermark_image:
                st.error("Please upload a watermark image")
            else:
                with st.spinner("Adding watermark..."):
                    uploaded_file.seek(0)
                    watermarked_pdf = add_watermark_simple(
                        uploaded_file,
                        watermark_text=watermark_text if watermark_type == "Text" else None,
                        watermark_image=watermark_image if watermark_type == "Image" else None,
                        opacity=opacity,
                        position=position,
                        font_size=font_size,
                        angle=angle,
                        scale=scale
                    )
                    
                    if watermarked_pdf:
                        st.success("✅ Watermark added successfully!")
                        
                        st.download_button(
                            label="📥 Download Watermarked PDF",
                            data=watermarked_pdf,
                            file_name="watermarked_document.pdf",
                            mime="application/pdf",
                            use_container_width=True
                        )
                        
                        show_preview_options(watermarked_pdf, "Watermarked PDF Result")


elif tool == "✏️ Edit PDF":
    st.header("✏️ Edit PDF")
    st.write("Edit your PDF by removing, keeping, or reordering pages.")
    
    uploaded_file = st.file_uploader("Choose a PDF file", type=['pdf'], key="edit")
    
    if uploaded_file:
        reader = PdfReader(uploaded_file)
        total_pages = len(reader.pages)
        
        col1, col2 = st.columns(2)
        with col1:
            st.metric("Total Pages", total_pages)
        
        show_preview_options(uploaded_file, "Original PDF (see page numbers)")
        
        edit_option = st.radio(
            "Edit option:",
            ["Remove pages", "Keep specific pages", "Reorder pages"],
            horizontal=True
        )
        
        if edit_option == "Remove pages":
            st.info("💡 Enter the page numbers you want to remove")
            pages_input = st.text_input(
                "Pages to remove (comma-separated):",
                placeholder="2, 4, 6"
            )
            
            if st.button("🗑️ Remove Pages", type="primary", use_container_width=True) and pages_input:
                try:
                    pages_to_remove = [int(p.strip()) for p in pages_input.split(',')]
                    pages_to_keep = [i for i in range(1, total_pages + 1) if i not in pages_to_remove]
                    
                    with st.spinner("Removing pages..."):
                        uploaded_file.seek(0)
                        edited_pdf = extract_pages(uploaded_file, pages_to_keep)
                        
                        if edited_pdf:
                            st.success(f"✅ Removed {len(pages_to_remove)} page(s)!")
                            
                            st.download_button(
                                label="📥 Download Edited PDF",
                                data=edited_pdf,
                                file_name="edited_document.pdf",
                                mime="application/pdf",
                                use_container_width=True
                            )
                            
                            show_preview_options(edited_pdf, "Edited PDF Result")
                except Exception as e:
                    st.error(f"Error: {str(e)}")
        
        elif edit_option == "Keep specific pages":
            st.info("💡 Only the pages you specify will be kept")
            pages_input = st.text_input(
                "Pages to keep (comma-separated):",
                placeholder="1, 3, 5, 7"
            )
            
            if st.button("✅ Keep Pages", type="primary", use_container_width=True) and pages_input:
                try:
                    pages_to_keep = [int(p.strip()) for p in pages_input.split(',')]
                    
                    with st.spinner("Processing..."):
                        uploaded_file.seek(0)
                        edited_pdf = extract_pages(uploaded_file, pages_to_keep)
                        
                        if edited_pdf:
                            st.success(f"✅ Kept {len(pages_to_keep)} page(s)!")
                            
                            st.download_button(
                                label="📥 Download Edited PDF",
                                data=edited_pdf,
                                file_name="edited_document.pdf",
                                mime="application/pdf",
                                use_container_width=True
                            )
                            
                            show_preview_options(edited_pdf, "Edited PDF Result")
                except Exception as e:
                    st.error(f"Error: {str(e)}")
        
        else:
            st.info("💡 Enter page numbers in the order you want them (e.g., 3, 1, 2 → page 3 becomes first)")
            pages_input = st.text_input(
                "New page order (comma-separated):",
                placeholder="3, 1, 2, 4"
            )
            
            if st.button("🔀 Reorder Pages", type="primary", use_container_width=True) and pages_input:
                try:
                    new_order = [int(p.strip()) for p in pages_input.split(',')]
                    
                    with st.spinner("Reordering pages..."):
                        uploaded_file.seek(0)
                        edited_pdf = extract_pages(uploaded_file, new_order)
                        
                        if edited_pdf:
                            st.success("✅ Pages reordered successfully!")
                            
                            st.download_button(
                                label="📥 Download Edited PDF",
                                data=edited_pdf,
                                file_name="reordered_document.pdf",
                                mime="application/pdf",
                                use_container_width=True
                            )
                            
                            show_preview_options(edited_pdf, "Reordered PDF Result")
                except Exception as e:
                    st.error(f"Error: {str(e)}")


# Footer
st.markdown("---")
st.markdown(
    """
    <div style='text-align: center; padding: 2rem;'>
        <h3 style='color: #00ff88;'>📄 PDF Toolkit</h3>
        <p style='color: #d0d0d0;'>Built with ❤️ using Python</p>
        <p style='font-size: 0.9rem; color: #999;'>All processing is done locally - your files never leave your device</p>
    </div>
    """,
    unsafe_allow_html=True
)
